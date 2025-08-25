from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Dict, List, Tuple

import yaml
from docx import Document


def load_domain(domain_path: Path) -> dict:
    if not domain_path.exists():
        raise FileNotFoundError(f"Domain file not found: {domain_path}")
    with domain_path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def save_domain(domain_path: Path, domain: dict) -> None:
    with domain_path.open("w", encoding="utf-8") as f:
        yaml.safe_dump(domain, f, sort_keys=False, allow_unicode=True)


def normalize_intent(name: str) -> str:
    return re.sub(r"\s+", "_", name.strip()).lower()


def parse_docx(docx_path: Path) -> List[Tuple[str, str]]:
    """
    Parse the DOCX for intent-response pairs.
    Supported formats:
      1) Tables with columns [intent, response] (order-insensitive, exact names preferred)
      2) Headings/paragraphs like:
           Intent: some_intent
           Response: the text to utter
    Returns: list of (intent, response)
    """
    document = Document(str(docx_path))
    pairs: List[Tuple[str, str]] = []

    # Strategy 1: parse tables
    for table in document.tables:
        headers = [normalize_intent(c.text) for c in table.rows[0].cells]
        if len(headers) < 2:
            continue
        # Try to find columns
        intent_idx = None
        response_idx = None
        for i, h in enumerate(headers):
            if "intent" in h:
                intent_idx = i
            if "response" in h or "utter" in h:
                response_idx = i
        if intent_idx is None or response_idx is None:
            continue

        for row in table.rows[1:]:
            try:
                intent = normalize_intent(row.cells[intent_idx].text)
                response = row.cells[response_idx].text.strip()
            except Exception:
                continue
            if intent and response:
                pairs.append((intent, response))

    if pairs:
        return pairs

    # Strategy 2: parse paragraphs with markers
    current_intent: str | None = None
    buffer: List[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m_intent = re.match(r"(?i)intent\s*[:：]\s*(.+)$", text)
        m_response = re.match(r"(?i)response\s*[:：]\s*(.+)$", text)
        if m_intent:
            # flush previous
            if current_intent and buffer:
                pairs.append((current_intent, "\n".join(buffer).strip()))
                buffer = []
            current_intent = normalize_intent(m_intent.group(1))
            continue
        if m_response:
            # explicit response line
            resp = m_response.group(1).strip()
            if current_intent and resp:
                pairs.append((current_intent, resp))
                current_intent = None
                buffer = []
            continue
        # Accumulate free text under current intent
        if current_intent:
            buffer.append(text)
    if current_intent and buffer:
        pairs.append((current_intent, "\n".join(buffer).strip()))

    return pairs


def upsert_domain_with_pairs(domain: dict, pairs: List[Tuple[str, str]], add_rules: bool) -> dict:
    # Ensure keys
    domain.setdefault("version", "3.1")
    domain.setdefault("intents", [])
    domain.setdefault("responses", {})

    # Build mapping: intent -> utter_action
    for intent, response in pairs:
        if intent not in domain["intents"]:
            domain["intents"].append(intent)

        utter_action = f"utter_{intent}"
        existing: List[Dict[str, str]] = domain["responses"].get(utter_action, [])
        # avoid duplicate exact text
        if not any(r.get("text") == response for r in existing):
            existing.append({"text": response})
        domain["responses"][utter_action] = existing

    if add_rules:
        # We do not modify domain for rules; they live in data/rules.yml
        pass
    return domain


def upsert_rules(rules_path: Path, pairs: List[Tuple[str, str]]) -> None:
    rules: dict = {}
    if rules_path.exists():
        with rules_path.open("r", encoding="utf-8") as f:
            rules = yaml.safe_load(f) or {}
    rules.setdefault("version", "3.1")
    rules.setdefault("rules", [])

    existing_rules = rules["rules"]
    existing_keys = {
        (tuple(step.items()) for step in r.get("steps", []))
        for r in existing_rules
    }

    for intent, _ in pairs:
        rule_name = f"Respond to {intent}"
        steps = [
            {"intent": intent},
            {"action": f"utter_{intent}"},
        ]
        # naive duplicate check
        if any(r.get("rule") == rule_name for r in existing_rules):
            continue
        existing_rules.append({"rule": rule_name, "steps": steps})

    with rules_path.open("w", encoding="utf-8") as f:
        yaml.safe_dump(rules, f, sort_keys=False, allow_unicode=True)


def main() -> None:
    parser = argparse.ArgumentParser(description="Import intent->response from DOCX into Rasa domain and rules")
    parser.add_argument("docx", type=Path, help="Path to the DOCX file")
    parser.add_argument("--domain", type=Path, default=Path("rasa/domain.yml"))
    parser.add_argument("--rules", type=Path, default=Path("rasa/data/rules.yml"))
    parser.add_argument("--no-rules", action="store_true", help="Do not write rules; only responses")
    args = parser.parse_args()

    pairs = parse_docx(args.docx)
    if not pairs:
        raise SystemExit("No intent-response pairs found in DOCX. Ensure it has a table (Intent/Response) or 'Intent:'/'Response:' sections.")

    domain = load_domain(args.domain)
    domain = upsert_domain_with_pairs(domain, pairs, add_rules=not args.no_rules)
    save_domain(args.domain, domain)

    if not args.no_rules:
        upsert_rules(args.rules, pairs)

    print(f"Imported {len(pairs)} intent-response pairs into {args.domain}")
    if not args.no_rules:
        print(f"Rules updated at {args.rules}")


if __name__ == "__main__":
    main()


